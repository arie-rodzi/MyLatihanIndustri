import streamlit as st
import sqlite3
from docx import Document
import os
from datetime import date
import base64

# Semakan login
if st.session_state.get("user_role") != "pelajar":
    st.warning("Modul ini hanya untuk pelajar.")
    st.stop()

st.set_page_config(page_title="Modul 2: Cetak Borang Permohonan", layout="wide")
st.title("📄 Modul 2: Cetak Borang Permohonan Latihan Industri")

pelajar_id = st.session_state.get("user_id", "")

# Sambung ke database
conn = sqlite3.connect("database/latihan_industri.db")
c = conn.cursor()

# Ambil maklumat pelajar
c.execute("SELECT nama, no_ic, kod_program FROM maklumat_pelajar WHERE pelajar_id=?", (pelajar_id,))
pelajar = c.fetchone()

# Ambil maklumat kelulusan
c.execute("SELECT status_lulus, nama_penyelaras, email_penyelaras, kod_program, tarikh_lulus FROM status_permohonan WHERE pelajar_id=?", (pelajar_id,))
status = c.fetchone()

if not pelajar:
    st.warning("Maklumat pelajar belum lengkap. Sila isi Modul 1 terlebih dahulu.")
    st.stop()

if not status or status[0].lower() != "lulus":
    st.info("Permohonan anda belum diluluskan oleh penyelaras program.")
    st.stop()

# Dapatkan data
nama, ic, program = pelajar
email = ""  # fallback emel pelajar jika tiada
_, nama_penyelaras, email_penyelaras, kod_program, _ = status
today = date.today().strftime("%Y-%m-%d")

# Folder simpan fail
folder_path = os.path.join("generated", program, "permohonan")
os.makedirs(folder_path, exist_ok=True)
docx_path = os.path.join(folder_path, f"permohonan_{pelajar_id}.docx")
pdf_path = os.path.join(folder_path, f"permohonan_{pelajar_id}.pdf")

# Gantian template
template_path = "templates/NS_SLI01_DLI01_BLI02.FIXED.docx"
doc = Document(template_path)

for p in doc.paragraphs:
    p.text = p.text.replace("«NOMBOR_ID_PELAJAR»", pelajar_id)
    p.text = p.text.replace("«NOMBOR_KAD_PENGENALAN»", ic)
    p.text = p.text.replace("«NAMA_PENUH_HURUF_BESAR»", nama.upper())
    p.text = p.text.replace("«NAMA_PROGRAM»", program)
    p.text = p.text.replace("«TARIKH_SURAT»", today)
    p.text = p.text.replace("«TARIKH_MULA_LI»", "2025-09-02")
    p.text = p.text.replace("«TARIKH_TAMAT_LI»", "2025-12-20")
    p.text = p.text.replace("«NAMA_PENYELARAS»", nama_penyelaras)
    p.text = p.text.replace("«email_penyelaras»", email_penyelaras)
    p.text = p.text.replace("«kod_pogram»", kod_program)
    p.text = p.text.replace("«ALAMAT_EMEL»", email)
    p.text = p.text.replace("«EMEL_PELAJAR»", email)

# Simpan fail DOCX
doc.save(docx_path)

# Papar isi surat sebagai teks
st.subheader("📝 Pratonton Kandungan Surat Permohonan")
doc_preview = Document(docx_path)
for para in doc_preview.paragraphs:
    st.write(para.text)

# --- MUAT NAIK PDF MANUAL ---
st.markdown("### 📤 Muat Naik Surat Permohonan (PDF)")
uploaded_pdf = st.file_uploader("Sila muat naik fail PDF surat permohonan:", type=["pdf"])

if uploaded_pdf:
    with open(pdf_path, "wb") as f:
        f.write(uploaded_pdf.read())
    st.success("✅ PDF berjaya dimuat naik.")

# --- PAPAR PDF JIKA SUDAH DIMUAT NAIK ---
if os.path.exists(pdf_path):
    with open(pdf_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode("utf-8")
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px"></iframe>'
        st.markdown("### 📄 Pratonton PDF Dimuat Naik")
        st.markdown(pdf_display, unsafe_allow_html=True)

        # Butang muat turun semula
        st.download_button("📥 Muat Turun Surat (PDF)", f, file_name=f"Surat_Permohonan_{pelajar_id}.pdf")

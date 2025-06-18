import streamlit as st
import os
import sqlite3
from docx import Document

# Konfigurasi
st.set_page_config(page_title="Cetak Surat Penempatan", layout="wide")
st.title("📄 Cetak Surat Penempatan")

if st.session_state.get("user_role") != "pelajar":
    st.warning("Modul ini hanya untuk pelajar.")
    st.stop()

pelajar_id = st.session_state.get("user_id", "")
docx_template = "templates/Surat_Penempatan_Final.docx"
output_docx = f"generated/surat_penempatan_{pelajar_id}.docx"
output_pdf = f"generated/surat_penempatan_{pelajar_id}.pdf"
os.makedirs("generated", exist_ok=True)

# Sambung ke pangkalan data
conn = sqlite3.connect("database/latihan_industri.db")
c = conn.cursor()

# Ambil data pelajar
c.execute("SELECT * FROM maklumat_pelajar WHERE pelajar_id=?", (pelajar_id,))
pelajar = c.fetchone()

# Ambil data industri
c.execute("SELECT * FROM maklumat_industri WHERE pelajar_id=?", (pelajar_id,))
industri = c.fetchone()

# Ambil data penyelaras
c.execute("SELECT * FROM penyelaras WHERE kod_program=?", (pelajar[3],))
penyelaras = c.fetchone()

if not pelajar or not industri or not penyelaras:
    st.warning("Sila lengkapkan maklumat pelajar, industri dan penyelaras dahulu.")
    st.stop()

# Data gantian
gantian = {
    "<<NAMA_PELAJAR>>": pelajar[1],
    "<<NO_KAD_PENGENALAN>>": pelajar[2],
    "<<NO_PELAJAR>>": pelajar[0],
    "<<NAMA_PROGRAM>>": penyelaras[1],
    "<<NAMA_ORGANISASI>>": industri[1],
    "<<ALAMAT_ORGANISASI_1>>": industri[2],
    "<<ALAMAT_ORGANISASI_2>>": industri[3],
    "<<BANDAR>>": industri[4],
    "<<POSKOD>>": industri[5],
    "<<NEGERI>>": industri[6],
    "<<NAMA_PEGAWAI>>": industri[7],
    "<<TARIKH_MULA_LATIHAN_INDUSTRI>>": industri[8],
    "<<TARIKH_TAMAT_LATIHAN_INDUSTRI>>": industri[9],
    "<<NAMA_PENYELARAS>>": penyelaras[2],
    "<<JAWATAN_PENYELARAS>>": penyelaras[3],
    "<<TELEFON_PENYELARAS>>": penyelaras[4],
    "<<EMEL_PENYELARAS>>": penyelaras[5],
}

def replace_tags(doc, replacements):
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                for run in p.runs:
                    run.text = run.text.replace(key, str(val))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(val))

if st.button("📄 Papar Surat & Sedia Untuk Cetak"):
    try:
        # Ganti tag dalam template dan simpan sebagai .docx
        doc = Document(docx_template)
        replace_tags(doc, gantian)
        doc.save(output_docx)

        # Paparkan isi kandungan surat sebagai teks
        st.write("### Pratonton Kandungan Surat (Teks Sahaja)")
        preview_doc = Document(output_docx)
        for para in preview_doc.paragraphs:
            st.write(para.text)

        # Beri muat turun fail PDF (jika telah disediakan oleh admin secara manual)
        if os.path.exists(output_pdf):
            with open(output_pdf, "rb") as f:
                st.download_button("📥 Muat Turun Surat (PDF)", f, file_name=f"Surat_Penempatan_{pelajar_id}.pdf")
        else:
            st.info("📎 PDF belum dijana. Sila cetak fail .docx secara manual dan simpan sebagai PDF.")

    except Exception as e:
        st.error(f"❌ Ralat semasa paparan surat: {e}")
